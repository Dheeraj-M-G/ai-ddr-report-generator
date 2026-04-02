"""
Build Detailed Diagnostic Report (DDR) as DOCX and PDF from structured JSON.
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from PIL import Image as PILImage
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from utils import EXTRACTED_IMAGES_DIR, OUTPUTS_DIR, ensure_directories, get_logger

logger = get_logger(__name__)


def _resolve_image_path(image_reference: str) -> Optional[Path]:
    """Map LLM image_reference basename to a file under extracted_images."""
    if not image_reference or image_reference.strip() == "Not Available":
        return None
    name = Path(image_reference.strip()).name
    p = EXTRACTED_IMAGES_DIR / name
    if p.is_file():
        return p
    for f in EXTRACTED_IMAGES_DIR.glob("*"):
        if f.name == name:
            return f
    return None


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _add_para(doc: Document, text: str, bullet: bool = False) -> None:
    if bullet:
        doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph(text)


def _escape_rl(text: str) -> str:
    text = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def build_ddr_docx(data: Dict[str, Any], stem: str = "DDR_Report") -> Tuple[Path, Path]:
    """
    Create DOCX and PDF under outputs/.

    Returns (docx_path, pdf_path).
    """
    ensure_directories()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = OUTPUTS_DIR / f"{stem}_{timestamp}.docx"
    pdf_path = OUTPUTS_DIR / f"{stem}_{timestamp}.pdf"

    doc = Document()
    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.alignment = 1
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    _add_heading(doc, "1. Property Issue Summary", level=1)
    _add_para(doc, str(data.get("property_issue_summary") or "Not Available"))

    _add_heading(doc, "2. Area-wise Observations", level=1)
    observations: List[Dict[str, Any]] = data.get("observations") or []
    if not observations:
        _add_para(doc, "No structured observations were produced. See section 7.")
    for i, obs in enumerate(observations, start=1):
        area = obs.get("area") or "Not Available"
        _add_heading(doc, f"2.{i} {area}", level=2)
        pairs = [
            ("Issue", obs.get("issue")),
            ("Description", obs.get("description")),
            ("Thermal observation", obs.get("thermal_observation")),
            ("Combined insight", obs.get("combined_insight")),
            ("Severity", obs.get("severity")),
            ("Recommendation", obs.get("recommendation")),
        ]
        for label, val in pairs:
            v = val if val is not None else "Not Available"
            s = str(v).strip() if v is not None else "Not Available"
            _add_para(doc, f"{label}: {s}")

        img_ref = obs.get("image_reference") or "Not Available"
        img_path = _resolve_image_path(str(img_ref))
        if img_path and img_path.is_file():
            try:
                doc.add_paragraph("Related image:")
                doc.add_picture(str(img_path), width=Inches(5.5))
            except Exception as e:
                logger.warning("Could not embed image %s: %s", img_path, e)
                _add_para(doc, "Image Not Available")
        else:
            _add_para(doc, "Image Not Available")
        doc.add_paragraph("")

    conflicts = data.get("conflicts") or []
    if conflicts:
        _add_heading(doc, "Conflicts between sources (see also notes)", level=2)
        for c in conflicts:
            topic = c.get("topic") or "Not Available"
            ins = c.get("inspection_says") or "Not Available"
            th = c.get("thermal_says") or "Not Available"
            _add_para(doc, f"Topic: {topic}", bullet=True)
            _add_para(doc, f"Inspection: {ins}", bullet=True)
            _add_para(doc, f"Thermal: {th}", bullet=True)
        doc.add_paragraph("")

    _add_heading(doc, "3. Probable Root Cause", level=1)
    _add_para(doc, str(data.get("probable_root_cause") or "Not Available"))

    _add_heading(doc, "4. Severity Assessment (with reasoning)", level=1)
    sev = data.get("severity_assessment") or {}
    if isinstance(sev, dict):
        _add_para(doc, f"Overall: {sev.get('overall') or 'Not Available'}")
        _add_para(doc, f"Reasoning: {sev.get('reasoning') or 'Not Available'}")
    else:
        _add_para(doc, "Not Available")

    _add_heading(doc, "5. Recommended Actions", level=1)
    actions = data.get("recommended_actions") or []
    if isinstance(actions, list) and actions:
        for a in actions:
            _add_para(doc, str(a), bullet=True)
    else:
        _add_para(doc, "Not Available")

    _add_heading(doc, "6. Additional Notes", level=1)
    _add_para(doc, str(data.get("additional_notes") or "Not Available"))

    _add_heading(doc, "7. Missing or Unclear Information", level=1)
    missing = data.get("missing_or_unclear") or []
    if isinstance(missing, list) and missing:
        for m in missing:
            _add_para(doc, str(m), bullet=True)
    else:
        _add_para(doc, "Not Available")

    doc.save(str(docx_path))
    logger.info("Wrote DOCX: %s", docx_path)

    _build_pdf_reportlab(data, pdf_path)
    return docx_path, pdf_path


def _build_pdf_reportlab(data: Dict[str, Any], pdf_path: Path) -> None:
    """Mirror DOCX sections in a simple PDF (ReportLab)."""
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=8,
        textColor=colors.HexColor("#1a1a1a"),
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=11,
        spaceAfter=6,
    )
    story: List = []

    story.append(Paragraph(_escape_rl("Detailed Diagnostic Report (DDR)"), h1))
    story.append(Spacer(1, 0.12 * inch))
    story.append(
        Paragraph(
            _escape_rl(datetime.now().strftime("Generated: %Y-%m-%d %H:%M")),
            body,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(_escape_rl("1. Property Issue Summary"), h1))
    story.append(
        Paragraph(_escape_rl(str(data.get("property_issue_summary") or "Not Available")), body)
    )
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph(_escape_rl("2. Area-wise Observations"), h1))
    observations: List[Dict[str, Any]] = data.get("observations") or []
    if not observations:
        story.append(Paragraph(_escape_rl("No structured observations were produced."), body))
    for i, obs in enumerate(observations, start=1):
        area = obs.get("area") or "Not Available"
        story.append(Paragraph(_escape_rl(f"2.{i} {area}"), h2))
        for label, key in [
            ("Issue", "issue"),
            ("Description", "description"),
            ("Thermal observation", "thermal_observation"),
            ("Combined insight", "combined_insight"),
            ("Severity", "severity"),
            ("Recommendation", "recommendation"),
        ]:
            val = obs.get(key) or "Not Available"
            story.append(Paragraph(_escape_rl(f"{label}: {val}"), body))
        img_ref = obs.get("image_reference") or "Not Available"
        img_path = _resolve_image_path(str(img_ref))
        if img_path and img_path.is_file():
            try:
                pil = PILImage.open(str(img_path))
                pw, ph = pil.size
                rw = 4.5 * inch
                rh = rw * (ph / float(pw)) if pw else 3 * inch
                story.append(RLImage(str(img_path), width=rw, height=rh))
            except Exception as e:
                logger.warning("PDF image skip %s: %s", img_path, e)
                story.append(Paragraph(_escape_rl("Image Not Available"), body))
        else:
            story.append(Paragraph(_escape_rl("Image Not Available"), body))
        story.append(Spacer(1, 0.1 * inch))

    conflicts = data.get("conflicts") or []
    if conflicts:
        story.append(Paragraph(_escape_rl("Conflicts between sources"), h2))
        for c in conflicts:
            topic = c.get("topic") or "Not Available"
            ins = c.get("inspection_says") or "Not Available"
            th = c.get("thermal_says") or "Not Available"
            story.append(Paragraph(_escape_rl(f"Topic: {topic}"), body))
            story.append(Paragraph(_escape_rl(f"Inspection: {ins}"), body))
            story.append(Paragraph(_escape_rl(f"Thermal: {th}"), body))
        story.append(Spacer(1, 0.1 * inch))

    story.append(Paragraph(_escape_rl("3. Probable Root Cause"), h1))
    story.append(Paragraph(_escape_rl(str(data.get("probable_root_cause") or "Not Available")), body))
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph(_escape_rl("4. Severity Assessment (with reasoning)"), h1))
    sev = data.get("severity_assessment") or {}
    if isinstance(sev, dict):
        story.append(
            Paragraph(
                _escape_rl(f"Overall: {sev.get('overall') or 'Not Available'}"),
                body,
            )
        )
        story.append(
            Paragraph(
                _escape_rl(f"Reasoning: {sev.get('reasoning') or 'Not Available'}"),
                body,
            )
        )
    else:
        story.append(Paragraph(_escape_rl("Not Available"), body))
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph(_escape_rl("5. Recommended Actions"), h1))
    actions = data.get("recommended_actions") or []
    if isinstance(actions, list) and actions:
        for a in actions:
            story.append(Paragraph(_escape_rl(f"• {a}"), body))
    else:
        story.append(Paragraph(_escape_rl("Not Available"), body))
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph(_escape_rl("6. Additional Notes"), h1))
    story.append(Paragraph(_escape_rl(str(data.get("additional_notes") or "Not Available")), body))
    story.append(Spacer(1, 0.12 * inch))

    story.append(Paragraph(_escape_rl("7. Missing or Unclear Information"), h1))
    missing = data.get("missing_or_unclear") or []
    if isinstance(missing, list) and missing:
        for m in missing:
            story.append(Paragraph(_escape_rl(f"• {m}"), body))
    else:
        story.append(Paragraph(_escape_rl("Not Available"), body))

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        rightMargin=inch * 0.85,
        leftMargin=inch * 0.85,
        topMargin=inch * 0.75,
        bottomMargin=inch * 0.75,
    )
    doc.build(story)
    logger.info("Wrote PDF: %s", pdf_path)


def save_json_snapshot(data: Dict[str, Any], stem: str = "DDR_structured") -> Path:
    """Save intermediate JSON next to reports for audit/debug."""
    ensure_directories()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    p = OUTPUTS_DIR / f"{stem}_{timestamp}.json"
    p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return p
